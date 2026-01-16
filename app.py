import os
import io
import re
import time
import random
from typing import Dict, Any, Optional

import streamlit as st
from PIL import Image

import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted, GoogleAPIError
import docx

# ======================================================
# ページ設定
# ======================================================
st.set_page_config(
    page_title="MAGI SYSTEM: ANALYTICAL PROTOCOL",
    page_icon="💠",
    layout="wide",
)

# ======================================================
# MAGI風 モダンUI CSS (FUIデザイン)
# ======================================================
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;500;700;900&family=Share+Tech+Mono&display=swap');

    /* 全体設定 */
    .stApp {
        background-color: #050505;
        background-image: 
            linear-gradient(rgba(0, 20, 40, 0.9), rgba(5, 5, 10, 0.95)),
            url("data:image/svg+xml,%3Csvg width='60' height='60' viewBox='0 0 60 60' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='none' fill-rule='evenodd'%3E%3Cg fill='%231a2639' fill-opacity='0.4'%3E%3Cpath d='M36 34v-4h-2v4h-4v2h4v4h2v-4h4v-2h-4zm0-30V0h-2v4h-4v2h4v4h2V6h4V4h-4zM6 34v-4H4v4H0v2h4v4h2v-4h4v-2H6zM6 4V0H4v4H0v2h4v4h2V6h4V4H6z'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
        color: #d0f0ff;
        font-family: 'Share Tech Mono', monospace;
    }

    /* ヘッダー */
    .magi-header {
        border-bottom: 2px solid #ff4d00;
        padding: 20px 0;
        margin-bottom: 30px;
        display: flex;
        justify-content: space-between;
        align-items: flex-end;
        background: linear-gradient(90deg, rgba(255,77,0,0.1) 0%, rgba(0,0,0,0) 80%);
    }
    .magi-title {
        font-family: 'Orbitron', sans-serif;
        font-size: 42px;
        font-weight: 900;
        color: #ff4d00;
        letter-spacing: 0.15em;
        text-shadow: 0 0 10px rgba(255, 77, 0, 0.6);
        line-height: 1;
    }
    .magi-subtitle {
        font-size: 14px;
        color: #ff8c00;
        letter-spacing: 0.3em;
        margin-top: 5px;
    }
    .magi-sys-status {
        text-align: right;
        font-size: 12px;
        color: #00ffcc;
    }

    /* カードデザイン共通 */
    .magi-card {
        background: rgba(10, 15, 20, 0.85);
        border: 1px solid #334455;
        border-radius: 4px;
        padding: 15px;
        margin-bottom: 15px;
        box-shadow: 0 0 15px rgba(0, 0, 0, 0.8);
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
    }
    .magi-card::before {
        content: "";
        position: absolute;
        top: 0; left: 0; width: 100%; height: 2px;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.5), transparent);
        opacity: 0.3;
    }

    /* 各エージェントの色分け */
    .agent-logic {
        border-color: #00ccff;
        box-shadow: 0 0 10px rgba(0, 204, 255, 0.15);
    }
    .agent-logic h4 { color: #00ccff; text-shadow: 0 0 5px rgba(0,204,255,0.5); }
    
    .agent-human {
        border-color: #ff9900;
        box-shadow: 0 0 10px rgba(255, 153, 0, 0.15);
    }
    .agent-human h4 { color: #ff9900; text-shadow: 0 0 5px rgba(255,153,0,0.5); }

    .agent-reality {
        border-color: #ff3366;
        box-shadow: 0 0 10px rgba(255, 51, 102, 0.15);
    }
    .agent-reality h4 { color: #ff3366; text-shadow: 0 0 5px rgba(255,51,102,0.5); }

    .agent-media { border-color: #aa00ff; }
    .agent-media h4 { color: #d066ff; }
    
    .agent-title {
        font-family: 'Orbitron', sans-serif;
        font-size: 18px;
        letter-spacing: 0.1em;
        margin-bottom: 10px;
        border-bottom: 1px solid rgba(255,255,255,0.1);
        padding-bottom: 5px;
        display: flex;
        justify-content: space-between;
    }

    /* 判定表示 */
    .decision-box {
        font-family: 'Orbitron', sans-serif;
        font-weight: 700;
        font-size: 24px;
        text-align: center;
        padding: 8px 0;
        margin: 10px 0;
        border: 1px solid;
        letter-spacing: 0.2em;
    }
    .decision-go { color: #00ff66; border-color: #00ff66; background: rgba(0,255,102,0.1); }
    .decision-nogo { color: #ff0033; border-color: #ff0033; background: rgba(255,0,51,0.1); }
    .decision-hold { color: #ffcc00; border-color: #ffcc00; background: rgba(255,204,0,0.1); }

    /* 統合結果 */
    .magi-aggregator {
        background: linear-gradient(180deg, rgba(20,30,50,0.9) 0%, rgba(5,10,20,0.95) 100%);
        border: 1px solid #4d5cff;
        border-left: 5px solid #4d5cff;
        padding: 20px;
        margin-top: 20px;
    }
    .section-label {
        font-family: 'Orbitron', sans-serif;
        font-size: 14px;
        color: #6677aa;
        letter-spacing: 0.2em;
        margin-bottom: 10px;
        display: block;
    }

    /* 入力エリア */
    .stTextArea textarea {
        background-color: rgba(0,0,0,0.3) !important;
        border: 1px solid #334455 !important;
        color: #00ffcc !important;
        font-family: 'Noto Sans JP', sans-serif;
    }
    
    /* ボタン */
    .stButton button {
        background: linear-gradient(45deg, #1a2a4a, #0d1a2f);
        border: 1px solid #00ccff;
        color: #00ccff;
        font-family: 'Orbitron', sans-serif;
        font-weight: bold;
        letter-spacing: 0.1em;
        transition: all 0.2s;
    }
    .stButton button:hover {
        background: #00ccff;
        color: #000;
        box-shadow: 0 0 15px #00ccff;
    }

    /* SWOT Chips */
    .swot-grid { display: flex; flex-wrap: wrap; gap: 5px; margin-top: 5px; }
    .swot-tag {
        font-size: 11px;
        padding: 2px 8px;
        border: 1px solid;
        border-radius: 0;
        background: rgba(0,0,0,0.4);
    }
    .swot-s { color: #81c784; border-color: #81c784; }
    .swot-w { color: #e57373; border-color: #e57373; }
    .swot-o { color: #64b5f6; border-color: #64b5f6; }
    .swot-t { color: #ffb74d; border-color: #ffb74d; }

    /* ユーティリティ */
    .divider-h {
        height: 1px;
        background: linear-gradient(90deg, transparent, #4d5cff, transparent);
        margin: 20px 0;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ======================================================
# ヘッダー表示
# ======================================================
st.markdown(
    """
    <div class="magi-header">
        <div>
            <div class="magi-title">MAGI SYSTEM</div>
            <div class="magi-subtitle">SUPER COMPUTER SYSTEM 3.0</div>
        </div>
        <div class="magi-sys-status">
            CODE: 771<br>
            PRIORITY: AAA<br>
            STATUS: <span style="color:#00ff00; animation: blink 1s infinite;">ONLINE</span>
        </div>
    </div>
    <style>@keyframes blink{0%{opacity:1}50%{opacity:0.3}100%{opacity:1}}</style>
    """,
    unsafe_allow_html=True,
)

# ======================================================
# Gemini API 設定
# ======================================================
api_key = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY"))

if not api_key:
    st.error("🚨 CRITICAL ERROR: API KEY NOT FOUND.")
    st.info("Set GEMINI_API_KEY in Streamlit secrets or environment variables.")
    st.stop()

genai.configure(api_key=api_key)

# ======================================================
# モデル設定 (セッション管理)
# ======================================================
if "gemini_model_name" not in st.session_state:
    # 安定性を優先してデフォルトを 1.5 Flash に変更
    st.session_state["gemini_model_name"] = "gemini-1.5-flash"

# サイドバー設定
st.sidebar.markdown(
    "<div style='font-family:Orbitron; font-size:18px; color:#ff4d00; margin-bottom:10px;'>SYSTEM CONFIG</div>",
    unsafe_allow_html=True,
)

MODEL_CHOICES = {
    "Gemini 1.5 Flash (Stable)": "gemini-1.5-flash",
    "Gemini 2.0 Flash (Preview)": "gemini-2.0-flash",
    "Gemini 1.5 Pro (High-Spec)": "gemini-1.5-pro",
}

selected_model_label = st.sidebar.selectbox(
    "PROCESSING CORE",
    list(MODEL_CHOICES.keys()),
    index=0
)
st.session_state["gemini_model_name"] = MODEL_CHOICES[selected_model_label]


def get_gemini_model():
    return genai.GenerativeModel(st.session_state["gemini_model_name"])


# ======================================================
# ヘルパー関数 (リトライロジック付き)
# ======================================================
def clean_text(text: str) -> str:
    if not text: return ""
    return text.replace("*", "").strip()

def generate_with_retry(model, content, max_retries=3):
    """
    429エラー(ResourceExhausted)が発生した場合、
    指数バックオフ (Exponential Backoff) で待機して再試行するラッパー関数
    """
    for attempt in range(max_retries):
        try:
            return model.generate_content(content)
        except ResourceExhausted as e:
            # クォータ制限の場合
            wait_time = (2 ** attempt) + random.uniform(0, 1) # 1秒, 2秒, 4秒...と待機時間を増やす
            if attempt < max_retries - 1:
                st.toast(f"⚠️ SYSTEM BUSY (429). RETRYING IN {wait_time:.1f}s...", icon="⏳")
                time.sleep(wait_time)
                continue
            else:
                # リトライ上限到達
                raise e
        except Exception as e:
            raise e

def analyze_media(file, mime_type: str, prompt: str) -> str:
    """画像や音声を解析する汎用関数（リトライ付き）"""
    model = get_gemini_model()
    try:
        if mime_type.startswith("image"):
            content = [prompt, Image.open(file)]
        else:
            content = [prompt, {"mime_type": mime_type, "data": file.getvalue()}]
        
        # リトライ付きで実行
        resp = generate_with_retry(model, content)
        return clean_text(resp.text)
    except ResourceExhausted:
        return "ERROR: 429 Quota Exceeded. (System Overload)"
    except Exception as e:
        return f"ERROR: {str(e)}"

# ======================================================
# MAGI ロジック
# ======================================================
def call_magi_core(context: Dict[str, Any], enable_swot: bool) -> str | None:
    model = get_gemini_model()
    
    # 役割定義
    system_prompt = """
あなたはスーパーコンピュータシステム「MAGI」です。
以下の3つの人格（エージェント）と、メディア解析担当、そして統合判断を行うメインプロセッサとして振る舞ってください。

【構成エージェント】
1. **Magi-Logic (Melchior)**: 
   - 科学者としての「自分」。冷徹、論理的、効率重視、最新技術への信頼。感情を排し、データと確率で判断する。
2. **Magi-Human (Balthasar)**: 
   - 母としての「自分」。倫理的、感情的、保護的。人間性、幸福、リスク回避、子供の将来を優先する。
3. **Magi-Reality (Casper)**: 
   - 女としての「自分」。現実的、政治的、直感的。現状維持、コスト、人間関係の機微、個人の欲望を重視する。

【タスク】
ユーザーの入力（質問・テキスト・メディア情報）に対し、上記3つの視点から議論し、それぞれ「可決(Go)」「否決(No-Go)」「保留(Hold)」を判定せよ。
あえて意見を対立させること。Logicが推奨してもHumanが倫理で止め、Realityがコストで渋るような構図が望ましい。

【出力フォーマット】
必ず以下の形式で出力すること。Markdownの装飾は最小限にせよ。

[SECTION:MAGI-LOGIC]
判定: (可決/否決/保留)
見解: (論理的視点からの120文字以内のコメント。断定的な口調)

[SECTION:MAGI-HUMAN]
判定: (可決/否決/保留)
見解: (人間的・倫理的視点からの120文字以内のコメント。丁寧だが心配性な口調)

[SECTION:MAGI-REALITY]
判定: (可決/否決/保留)
見解: (現実的・政治的視点からの120文字以内のコメント。シニカルまたは打算的な口調)

[SECTION:MAGI-MEDIA]
判定: (可決/否決/保留)
見解: (デザイン・印象・表現面からの120文字以内のコメント)

[SECTION:INTEGRATION]
結論: (承認/否決/条件付き承認 など簡潔に)
詳細: (3者の意見を統合した最終アドバイス。300文字以内)
"""

    if enable_swot:
        system_prompt += """
[SECTION:SWOT]
Strengths: (強みを5つ、読点で区切って列挙)
Weaknesses: (弱みを5つ、読点で区切って列挙)
Opportunities: (機会を5つ、読点で区切って列挙)
Threats: (脅威を5つ、読点で区切って列挙)
"""

    user_data = f"""
    QUERY: {context['user_question']}
    ADDITIONAL_TEXT: {context['text_input']}
    VISUAL_DATA: {context['image_description']}
    AUDIO_DATA: {context['audio_transcript']}
    """

    try:
        # リトライ付きで実行
        response = generate_with_retry(model, [system_prompt, user_data])
        return response.text
    except ResourceExhausted:
        return "SYSTEM FAILURE: 429 RESOURCE EXHAUSTED. Please switch models or wait a moment."
    except Exception as e:
        return f"SYSTEM FAILURE: {str(e)}"

# ======================================================
# 解析ロジック (テキスト処理)
# ======================================================
def parse_magi_output(text: str):
    sections = {}
    pattern = re.compile(r"\[SECTION:(.*?)\]")
    parts = pattern.split(text)
    
    for i in range(1, len(parts), 2):
        tag = parts[i].strip()
        content = parts[i+1].strip()
        
        data = {"decision": "保留", "summary": "", "raw": content}
        
        for line in content.split('\n'):
            if line.startswith("判定:"):
                val = line.split(":", 1)[1].strip()
                if "可決" in val: data["decision"] = "可決"
                elif "否決" in val: data["decision"] = "否決"
                else: data["decision"] = "保留"
            elif line.startswith("見解:") or line.startswith("詳細:"):
                data["summary"] = line.split(":", 1)[1].strip()
        
        if tag == "SWOT":
            swot_data = {}
            for line in content.split('\n'):
                if ":" in line:
                    k, v = line.split(":", 1)
                    swot_data[k.strip()] = v.strip()
            sections["SWOT"] = swot_data
        else:
            sections[tag] = data

    return sections

def get_decision_style(decision):
    if decision == "可決": return "decision-go", "GO"
    if decision == "否決": return "decision-nogo", "NO-GO"
    return "decision-hold", "HOLD"

# ======================================================
# Word レポート作成
# ======================================================
def create_docx(context, sections, image=None):
    doc = docx.Document()
    doc.add_heading('MAGI ANALYTICAL REPORT', 0)
    
    doc.add_heading('1. INPUT DATA', level=1)
    doc.add_paragraph(f"Query: {context['user_question']}")
    if context['text_input']: doc.add_paragraph(f"Text: {context['text_input']}")
    
    if image:
        img_stream = io.BytesIO()
        image.save(img_stream, format="PNG")
        img_stream.seek(0)
        doc.add_picture(img_stream, width=docx.shared.Inches(2.5))

    doc.add_heading('2. MAGI DELIBERATION', level=1)
    
    agent_map = {
        "MAGI-LOGIC": "MELCHIOR-1 (Logic)",
        "MAGI-HUMAN": "BALTHASAR-2 (Human)",
        "MAGI-REALITY": "CASPER-3 (Reality)",
        "MAGI-MEDIA": "MEDIA ANALYZER"
    }
    
    for key, name in agent_map.items():
        if key in sections:
            sec = sections[key]
            p = doc.add_paragraph()
            p.add_run(f"[{name}] ").bold = True
            p.add_run(f"Vote: {sec['decision']}\n")
            p.add_run(sec['summary'])

    doc.add_heading('3. FINAL INTEGRATION', level=1)
    if "INTEGRATION" in sections:
        doc.add_paragraph(sections["INTEGRATION"]["raw"])

    if "SWOT" in sections:
        doc.add_heading('4. SWOT ANALYSIS', level=1)
        for k, v in sections["SWOT"].items():
            doc.add_paragraph(f"{k}: {v}")
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ======================================================
# UI 構築
# ======================================================

# --- サイドバー入力 ---
input_mode = st.sidebar.radio("DATA INPUT SOURCE", ["File Upload", "Camera", "None"], index=0)
uploaded_file = None
if input_mode == "File Upload":
    uploaded_file = st.sidebar.file_uploader("ARCHIVE DATA", type=["png", "jpg", "jpeg", "wav", "mp3", "txt"])
elif input_mode == "Camera":
    uploaded_file = st.sidebar.camera_input("VISUAL SENSOR")

swot_mode = st.sidebar.checkbox("ACTIVATE SWOT MODULE", value=False)

# --- メインエリア ---
st.markdown('<span class="section-label">:: USER QUERY ::</span>', unsafe_allow_html=True)
user_question = st.text_area("ENTER YOUR DILEMMA", height=80, placeholder="例：このプロジェクトを進めるべきか？ 今の状況を分析してほしい。")
text_input = st.text_area("SUPPLEMENTARY DATA (OPTIONAL)", height=80)

# 解析用コンテキスト
context = {
    "user_question": user_question,
    "text_input": text_input,
    "image_description": "",
    "audio_transcript": ""
}
report_image = None

# メディア処理
if uploaded_file:
    mime = uploaded_file.type
    st.markdown('<span class="section-label">:: MEDIA DATA ::</span>', unsafe_allow_html=True)
    
    if mime.startswith("image"):
        image = Image.open(uploaded_file)
        report_image = image
        st.image(image, caption="VISUAL DATA ACQUIRED", width=300)
        with st.spinner("ANALYZING VISUAL PATTERNS..."):
            context["image_description"] = analyze_media(
                uploaded_file, mime, 
                "この画像に写っているものを客観的に、詳細に描写してください。感情的な印象も含めてください。"
            )
            
    elif mime.startswith("audio"):
        st.audio(uploaded_file)
        with st.spinner("DECODING AUDIO WAVEFORM..."):
            context["audio_transcript"] = analyze_media(
                uploaded_file, mime, 
                "この音声を日本語に書き起こしてください。"
            )

# --- 実行ボタン ---
st.markdown("<div style='margin-top:20px;'></div>", unsafe_allow_html=True)
if st.button("INITIALIZE MAGI DELIBERATION", type="primary", use_container_width=True):
    
    if not user_question and not uploaded_file and not text_input:
        st.warning("⚠️ DATA INSUFFICIENT. PLEASE INPUT QUERY OR MEDIA.")
        st.stop()
        
    # プログレス演出
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    steps = [
        "CONNECTING TO MAGI SYSTEM...",
        "SYNCING WITH MELCHIOR-1...",
        "SYNCING WITH BALTHASAR-2...",
        "SYNCING WITH CASPER-3...",
        "CALCULATING PROBABILITIES...",
        "DELIBERATION IN PROGRESS..."
    ]
    
    for i, step in enumerate(steps):
        status_text.markdown(f"<span style='color:#00ffcc; font-family:Orbitron;'>{step}</span>", unsafe_allow_html=True)
        progress_bar.progress((i + 1) * 15)
        time.sleep(0.1) 

    # Gemini 実行
    raw_result = call_magi_core(context, swot_mode)
    progress_bar.progress(100)
    status_text.empty()
    progress_bar.empty()
    
    # 失敗時の表示
    if not raw_result or "SYSTEM FAILURE" in raw_result:
        st.error(raw_result or "UNKNOWN ERROR")
        if "RESOURCE EXHAUSTED" in raw_result:
             st.info("💡 **HINT**: Try switching to 'Gemini 1.5 Flash' in the sidebar or wait a minute before retrying.")
        st.stop()

    # 結果パース
    sections = parse_magi_output(raw_result)

    # ==================================================
    # 結果表示 (3カラムレイアウト)
    # ==================================================
    st.markdown('<div class="divider-h"></div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    # MELCHIOR (Logic)
    with col1:
        sec = sections.get("MAGI-LOGIC", {})
        style, label = get_decision_style(sec.get("decision"))
        st.markdown(f"""
        <div class="magi-card agent-logic">
            <div class="agent-title">
                <span>MELCHIOR-1</span>
                <span style="font-size:10px;">LOGIC</span>
            </div>
            <div class="decision-box {style}">{label}</div>
            <div style="font-size:13px; line-height:1.4;">{sec.get('summary', 'No Data')}</div>
        </div>
        """, unsafe_allow_html=True)

    # BALTHASAR (Human)
    with col2:
        sec = sections.get("MAGI-HUMAN", {})
        style, label = get_decision_style(sec.get("decision"))
        st.markdown(f"""
        <div class="magi-card agent-human">
            <div class="agent-title">
                <span>BALTHASAR-2</span>
                <span style="font-size:10px;">HUMAN</span>
            </div>
            <div class="decision-box {style}">{label}</div>
            <div style="font-size:13px; line-height:1.4;">{sec.get('summary', 'No Data')}</div>
        </div>
        """, unsafe_allow_html=True)

    # CASPER (Reality)
    with col3:
        sec = sections.get("MAGI-REALITY", {})
        style, label = get_decision_style(sec.get("decision"))
        st.markdown(f"""
        <div class="magi-card agent-reality">
            <div class="agent-title">
                <span>CASPER-3</span>
                <span style="font-size:10px;">REALITY</span>
            </div>
            <div class="decision-box {style}">{label}</div>
            <div style="font-size:13px; line-height:1.4;">{sec.get('summary', 'No Data')}</div>
        </div>
        """, unsafe_allow_html=True)

    # Media & Integration (下部)
    c_media, c_integ = st.columns([1, 2])
    
    with c_media:
        sec = sections.get("MAGI-MEDIA", {})
        style, label = get_decision_style(sec.get("decision"))
        st.markdown(f"""
        <div class="magi-card agent-media">
            <div class="agent-title">
                <span>MEDIA.OP</span>
                <span style="font-size:10px;">ARTS</span>
            </div>
            <div class="decision-box {style}">{label}</div>
            <div style="font-size:12px;">{sec.get('summary', 'No Data')}</div>
        </div>
        """, unsafe_allow_html=True)
        
    with c_integ:
        sec = sections.get("INTEGRATION", {})
        st.markdown(f"""
        <div class="magi-aggregator">
            <div class="agent-title" style="border:none; color:#fff;">:: FINAL DECISION ::</div>
            <div style="font-size:16px; margin-bottom:10px; color:#4d5cff; font-weight:bold;">
                {sec.get('raw', '').split('詳細:')[0].replace('結論:', '')}
            </div>
            <div style="font-size:14px; line-height:1.6; color:#d0f0ff;">
                {sec.get('summary', '')}
            </div>
        </div>
        """, unsafe_allow_html=True)

    # SWOT Module
    if swot_mode and "SWOT" in sections:
        swot = sections["SWOT"]
        st.markdown('<div class="divider-h"></div>', unsafe_allow_html=True)
        st.markdown('<span class="section-label">:: SWOT STRATEGIC GRID ::</span>', unsafe_allow_html=True)
        
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**STRENGTHS**")
            html = '<div class="swot-grid">' + "".join([f'<span class="swot-tag swot-s">{x}</span>' for x in swot.get('Strengths','').split('、')]) + '</div>'
            st.markdown(html, unsafe_allow_html=True)
            
            st.markdown("<br>**OPPORTUNITIES**", unsafe_allow_html=True)
            html = '<div class="swot-grid">' + "".join([f'<span class="swot-tag swot-o">{x}</span>' for x in swot.get('Opportunities','').split('、')]) + '</div>'
            st.markdown(html, unsafe_allow_html=True)
            
        with c2:
            st.markdown("**WEAKNESSES**")
            html = '<div class="swot-grid">' + "".join([f'<span class="swot-tag swot-w">{x}</span>' for x in swot.get('Weaknesses','').split('、')]) + '</div>'
            st.markdown(html, unsafe_allow_html=True)
            
            st.markdown("<br>**THREATS**", unsafe_allow_html=True)
            html = '<div class="swot-grid">' + "".join([f'<span class="swot-tag swot-t">{x}</span>' for x in swot.get('Threats','').split('、')]) + '</div>'
            st.markdown(html, unsafe_allow_html=True)

    # レポート出力
    docx_bytes = create_docx(context, sections, report_image)
    st.markdown('<div class="divider-h"></div>', unsafe_allow_html=True)
    st.download_button(
        label="💾 EXPORT REPORT (.DOCX)",
        data=docx_bytes,
        file_name="MAGI_CONFIDENTIAL_REPORT.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="secondary"
    )
